"""
课程导出服务
支持PDF、DOCX、HTML、JSON多格式实际文件生成
"""

import os
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.flowables import PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from jinja2 import Template


class CourseExportService:
    """课程导出服务类"""

    def __init__(self):
        self.export_dir = Path("/home/easegen/EduAgents/backend/exports")
        self.ensure_export_dirs()

    def ensure_export_dirs(self):
        """确保导出目录存在"""
        for format_type in ["pdf", "docx", "html", "json"]:
            (self.export_dir / format_type).mkdir(parents=True, exist_ok=True)

    async def export_course(self, course_data: Dict[str, Any], export_format: str,
                          include_resources: bool = True,
                          include_assessments: bool = True) -> Dict[str, Any]:
        """
        导出课程为指定格式

        Args:
            course_data: 课程数据
            export_format: 导出格式 (pdf, docx, html, json)
            include_resources: 是否包含资源
            include_assessments: 是否包含评估

        Returns:
            导出结果信息
        """

        export_id = str(uuid.uuid4())
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        course_id = course_data.get("course_id", "unknown")

        # 生成文件名
        filename = f"pbl_course_{course_id}_{timestamp}.{export_format}"
        file_path = self.export_dir / export_format / filename

        try:
            # 根据格式调用对应的导出方法
            if export_format == "pdf":
                actual_path = await self.export_to_pdf(course_data, file_path,
                                                     include_resources, include_assessments)
            elif export_format == "docx":
                actual_path = await self.export_to_docx(course_data, file_path,
                                                      include_resources, include_assessments)
            elif export_format == "html":
                actual_path = await self.export_to_html(course_data, file_path,
                                                      include_resources, include_assessments)
            elif export_format == "json":
                actual_path = await self.export_to_json(course_data, file_path,
                                                      include_resources, include_assessments)
            else:
                raise ValueError(f"不支持的导出格式: {export_format}")

            # 获取文件大小
            file_size = self.get_file_size(actual_path)

            return {
                "export_id": export_id,
                "course_id": course_id,
                "format": export_format,
                "file_name": filename,
                "file_path": str(actual_path),
                "file_size": file_size,
                "download_url": f"/api/v1/courses/download/{export_format}/{filename}",
                "includes": {
                    "course_outline": True,
                    "learning_activities": True,
                    "assessment_rubrics": include_assessments,
                    "teaching_resources": include_resources,
                    "student_materials": True,
                    "teacher_guide": True
                },
                "created_at": datetime.now().isoformat(),
                "expires_at": "2025-10-21T18:00:00Z"
            }

        except Exception as e:
            raise Exception(f"导出失败: {str(e)}")

    async def export_to_pdf(self, course_data: Dict[str, Any], file_path: Path,
                          include_resources: bool, include_assessments: bool) -> Path:
        """导出为PDF格式"""

        doc = SimpleDocTemplate(str(file_path), pagesize=A4)
        styles = getSampleStyleSheet()
        story = []

        # 使用更可靠的中文字体解决方案
        font_name = 'Helvetica'  # 默认字体
        try:
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase.cidfonts import UnicodeCIDFont

            font_registered = False

            # 方案1: 尝试使用reportlab内置的Unicode CID字体
            try:
                # 注册中文Unicode CID字体（这个通常支持中文）
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                font_name = 'STSong-Light'
                font_registered = True
                print("✅ 成功注册STSong-Light CID字体")
            except Exception as e:
                print(f"⚠️ STSong-Light注册失败: {e}")

            # 方案2: 如果CID字体失败，尝试TTF字体
            if not font_registered:
                font_paths = [
                    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                    "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                ]

                for font_path in font_paths:
                    if Path(font_path).exists():
                        try:
                            pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                            font_name = 'ChineseFont'
                            font_registered = True
                            print(f"✅ 成功注册TTF字体: {font_path}")
                            break
                        except Exception as e:
                            print(f"⚠️ TTF字体注册失败: {e}")
                            continue

            if not font_registered:
                print("⚠️ 所有字体注册失败，使用Helvetica（中文可能显示为方块）")

        except Exception as e:
            print(f"❌ 字体注册过程异常: {e}")
            font_name = 'Helvetica'

        # 自定义样式支持中文
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # 居中
            fontName=font_name
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=12,
            fontName=font_name
        )

        heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=styles['Heading2'],
            fontSize=14,
            fontName=font_name
        )

        heading3_style = ParagraphStyle(
            'CustomHeading3',
            parent=styles['Heading3'],
            fontSize=12,
            fontName=font_name
        )

        # 使用安全的文本处理方式
        def safe_text(text):
            """安全处理文本，对于无法显示的字符使用替代方案"""
            if font_name == 'Helvetica':
                # 如果使用默认字体，提供英文标题
                if "AI原生PBL课程设计" in text:
                    return "AI-Native PBL Course Design"
                elif "课程基本信息" in text:
                    return "Course Information"
                elif "学习目标" in text:
                    return "Learning Objectives"
                elif "项目驱动问题" in text:
                    return "Driving Question"
                elif "最终产品" in text:
                    return "Final Products"
                elif "课程实施阶段" in text:
                    return "Implementation Phases"
                elif "评估体系" in text:
                    return "Assessment System"
                elif "教学资源" in text:
                    return "Teaching Resources"
                elif "技术要求" in text:
                    return "Technical Requirements"
                elif "教师准备" in text:
                    return "Teacher Preparation"
                elif "课程质量指标" in text:
                    return "Quality Metrics"
            return text

        story.append(Paragraph(safe_text("AI原生PBL课程设计"), title_style))
        story.append(Paragraph(safe_text(course_data.get('title', 'Untitled Course')), title_style))
        story.append(Spacer(1, 20))

        # 课程基本信息
        story.append(Paragraph(safe_text("课程基本信息"), heading2_style))
        info_data = [
            [safe_text("课程名称") if font_name != 'Helvetica' else "Course Name", course_data.get('title', '')],
            [safe_text("教育层级") if font_name != 'Helvetica' else "Education Level", course_data.get('education_level', '')],
            [safe_text("年级") if font_name != 'Helvetica' else "Grade Levels", str(course_data.get('grade_levels', []))],
            [safe_text("持续周数") if font_name != 'Helvetica' else "Duration (weeks)", f"{course_data.get('duration_weeks', 0)} weeks"],
            [safe_text("总课时") if font_name != 'Helvetica' else "Total Hours", f"{course_data.get('duration_hours', 0)} hours"],
            [safe_text("创建时间") if font_name != 'Helvetica' else "Created At", course_data.get('created_at', '')[:19]]
        ]

        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('FONTSIZE', (0, 1), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(info_table)
        story.append(Spacer(1, 20))

        # 学习目标
        story.append(Paragraph(safe_text("学习目标"), heading2_style))
        objectives = course_data.get('learning_objectives', [])
        for i, obj in enumerate(objectives, 1):
            story.append(Paragraph(f"{i}. {obj}", normal_style))
        story.append(Spacer(1, 15))

        # 驱动性问题
        story.append(Paragraph(safe_text("项目驱动问题"), heading2_style))
        story.append(Paragraph(course_data.get('driving_question', ''), normal_style))
        story.append(Spacer(1, 15))

        # 最终产品
        story.append(Paragraph(safe_text("最终产品"), heading2_style))
        final_products = course_data.get('final_products', [])
        for product in final_products:
            story.append(Paragraph(f"• {product}", normal_style))
        story.append(Spacer(1, 15))

        # 课程阶段
        story.append(Paragraph(safe_text("课程实施阶段"), heading2_style))
        phases = course_data.get('phases', [])
        for phase in phases:
            story.append(Paragraph(f"【{phase.get('name', '')}】 - {phase.get('duration', '')}", heading3_style))
            activities = phase.get('activities', [])
            for activity in activities:
                story.append(Paragraph(f"• {activity}", normal_style))

            # 添加AI工具信息
            ai_tools = phase.get('ai_tools', [])
            if ai_tools:
                ai_tools_text = f"Recommended AI Tools: {', '.join(ai_tools)}" if font_name == 'Helvetica' else f"推荐AI工具: {', '.join(ai_tools)}"
                story.append(Paragraph(ai_tools_text, normal_style))
            story.append(Spacer(1, 10))

        # 评估体系
        if include_assessments and 'assessments' in course_data:
            story.append(PageBreak())
            story.append(Paragraph(safe_text("评估体系"), heading2_style))
            assessments = course_data.get('assessments', [])
            for assessment in assessments:
                story.append(Paragraph(f"【{assessment.get('name', '')}】", heading3_style))
                story.append(Paragraph(f"类型: {assessment.get('type', '')}", normal_style))
                story.append(Paragraph(f"权重: {assessment.get('weight', 0)*100}%", normal_style))
                methods = assessment.get('methods', [])
                for method in methods:
                    story.append(Paragraph(f"• {method}", normal_style))
                story.append(Spacer(1, 10))

        # 教学资源
        if include_resources and 'resources' in course_data:
            story.append(Paragraph(safe_text("教学资源"), heading2_style))
            resources = course_data.get('resources', [])
            for resource in resources:
                story.append(Paragraph(f"【{resource.get('title', '')}】", heading3_style))
                story.append(Paragraph(f"类型: {resource.get('type', '')}", normal_style))
                story.append(Paragraph(f"描述: {resource.get('description', '')}", normal_style))
                story.append(Spacer(1, 10))

        # 技术要求
        if 'technology_requirements' in course_data:
            story.append(Paragraph("技术要求", heading2_style))
            tech_requirements = course_data.get('technology_requirements', [])
            for req in tech_requirements:
                story.append(Paragraph(f"• {req}", normal_style))
            story.append(Spacer(1, 10))

        # 教师准备
        if 'teacher_preparation' in course_data:
            story.append(Paragraph("教师准备", heading2_style))
            preparations = course_data.get('teacher_preparation', [])
            for prep in preparations:
                story.append(Paragraph(f"• {prep}", normal_style))
            story.append(Spacer(1, 10))

        # 质量指标（如果有）
        if 'quality_metrics' in course_data:
            story.append(Paragraph("课程质量指标", heading2_style))
            metrics = course_data.get('quality_metrics', {})
            story.append(Paragraph(f"AI能力覆盖度: {metrics.get('ai_competency_coverage', 0)*100:.0f}%", normal_style))
            story.append(Paragraph(f"PBL方法论完整性: {metrics.get('pbl_methodology_score', 0)*100:.0f}%", normal_style))
            story.append(Paragraph(f"内容丰富度: {metrics.get('content_richness', 0)*100:.0f}%", normal_style))
            story.append(Paragraph(f"评估真实性: {metrics.get('assessment_authenticity', 0)*100:.0f}%", normal_style))
            story.append(Paragraph(f"资源完整性: {metrics.get('resource_completeness', 0)*100:.0f}%", normal_style))

        # 生成PDF
        doc.build(story)
        return file_path

    async def export_to_docx(self, course_data: Dict[str, Any], file_path: Path,
                           include_resources: bool, include_assessments: bool) -> Path:
        """导出为DOCX格式"""

        doc = Document()

        # 标题
        title = doc.add_heading('AI原生PBL课程设计', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        course_title = doc.add_heading(course_data.get('title', '未命名课程'), 1)
        course_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 课程基本信息
        doc.add_heading('课程基本信息', 2)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        info_items = [
            ("课程名称", course_data.get('title', '')),
            ("教育层级", course_data.get('education_level', '')),
            ("年级", str(course_data.get('grade_levels', []))),
            ("持续周数", f"{course_data.get('duration_weeks', 0)}周"),
            ("总课时", f"{course_data.get('duration_hours', 0)}小时"),
            ("创建时间", course_data.get('created_at', '')[:19])
        ]

        for i, (key, value) in enumerate(info_items):
            table.cell(i, 0).text = key
            table.cell(i, 1).text = value

        # 学习目标
        doc.add_heading('学习目标', 2)
        objectives = course_data.get('learning_objectives', [])
        for i, obj in enumerate(objectives, 1):
            doc.add_paragraph(f"{i}. {obj}")

        # 驱动性问题
        doc.add_heading('项目驱动问题', 2)
        doc.add_paragraph(course_data.get('driving_question', ''))

        # 最终产品
        doc.add_heading('最终产品', 2)
        final_products = course_data.get('final_products', [])
        for product in final_products:
            doc.add_paragraph(f"• {product}")

        # 课程阶段
        doc.add_heading('课程实施阶段', 2)
        phases = course_data.get('phases', [])
        for phase in phases:
            doc.add_heading(f"{phase.get('name', '')} - {phase.get('duration', '')}", 3)
            activities = phase.get('activities', [])
            for activity in activities:
                doc.add_paragraph(f"• {activity}")

            # 添加AI工具信息
            ai_tools = phase.get('ai_tools', [])
            if ai_tools:
                doc.add_paragraph(f"推荐AI工具: {', '.join(ai_tools)}")
                doc.add_paragraph("")  # 空行

        # 评估体系
        if include_assessments and 'assessments' in course_data:
            doc.add_page_break()
            doc.add_heading('评估体系', 2)
            assessments = course_data.get('assessments', [])
            for assessment in assessments:
                doc.add_heading(assessment.get('name', ''), 3)
                doc.add_paragraph(f"类型: {assessment.get('type', '')}")
                doc.add_paragraph(f"权重: {assessment.get('weight', 0)*100}%")
                methods = assessment.get('methods', [])
                for method in methods:
                    doc.add_paragraph(f"• {method}")

        # 教学资源
        if include_resources and 'resources' in course_data:
            doc.add_heading('教学资源', 2)
            resources = course_data.get('resources', [])
            for resource in resources:
                doc.add_heading(resource.get('title', ''), 3)
                doc.add_paragraph(f"类型: {resource.get('type', '')}")
                doc.add_paragraph(f"描述: {resource.get('description', '')}")
                doc.add_paragraph("")  # 空行

        # 技术要求
        if 'technology_requirements' in course_data:
            doc.add_heading('技术要求', 2)
            tech_requirements = course_data.get('technology_requirements', [])
            for req in tech_requirements:
                doc.add_paragraph(f"• {req}")

        # 教师准备
        if 'teacher_preparation' in course_data:
            doc.add_heading('教师准备', 2)
            preparations = course_data.get('teacher_preparation', [])
            for prep in preparations:
                doc.add_paragraph(f"• {prep}")

        # 质量指标
        if 'quality_metrics' in course_data:
            doc.add_heading('课程质量指标', 2)
            metrics = course_data.get('quality_metrics', {})

            metrics_table = doc.add_table(rows=6, cols=2)
            metrics_table.style = 'Table Grid'

            metrics_items = [
                ("AI能力覆盖度", f"{metrics.get('ai_competency_coverage', 0)*100:.0f}%"),
                ("PBL方法论完整性", f"{metrics.get('pbl_methodology_score', 0)*100:.0f}%"),
                ("内容丰富度", f"{metrics.get('content_richness', 0)*100:.0f}%"),
                ("评估真实性", f"{metrics.get('assessment_authenticity', 0)*100:.0f}%"),
                ("资源完整性", f"{metrics.get('resource_completeness', 0)*100:.0f}%"),
                ("综合评分", f"{sum(metrics.values())/len(metrics)*5:.1f}/5.0")
            ]

            for i, (metric, value) in enumerate(metrics_items):
                metrics_table.cell(i, 0).text = metric
                metrics_table.cell(i, 1).text = value

        # 设计信息
        doc.add_heading('设计信息', 2)
        doc.add_paragraph(f"课程设计时间: {course_data.get('created_at', '')[:19]}")
        if 'design_agents' in course_data:
            doc.add_paragraph("参与设计的AI智能体:")
            agent_names = {
                'education_theorist': '教育理论专家',
                'course_architect': '课程架构师',
                'content_designer': '内容设计师',
                'assessment_expert': '评估专家',
                'material_creator': '素材创作者'
            }
            for agent in course_data.get('design_agents', []):
                doc.add_paragraph(f"• {agent_names.get(agent, agent)}")

        if course_data.get('ai_native'):
            doc.add_paragraph("✓ AI原生设计")
        if course_data.get('competency_based'):
            doc.add_paragraph("✓ 能力导向课程")

        # 保存文档
        doc.save(str(file_path))
        return file_path

    async def export_to_html(self, course_data: Dict[str, Any], file_path: Path,
                           include_resources: bool, include_assessments: bool) -> Path:
        """导出为HTML格式"""

        html_template = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ course.title }} - AI原生PBL课程设计</title>
    <style>
        body { font-family: "微软雅黑", Arial, sans-serif; margin: 40px; line-height: 1.6; }
        .header { text-align: center; margin-bottom: 40px; }
        .title { color: #2c3e50; font-size: 28px; margin-bottom: 10px; }
        .subtitle { color: #34495e; font-size: 20px; }
        .section { margin-bottom: 30px; }
        .section-title { color: #2980b9; font-size: 18px; border-bottom: 2px solid #3498db; padding-bottom: 5px; }
        .info-table { width: 100%; border-collapse: collapse; margin: 15px 0; }
        .info-table th, .info-table td { border: 1px solid #ddd; padding: 12px; text-align: left; }
        .info-table th { background-color: #f8f9fa; font-weight: bold; }
        .objective-list { list-style-type: decimal; padding-left: 20px; }
        .phase { margin: 20px 0; padding: 15px; background-color: #f8f9fa; border-radius: 5px; }
        .phase-title { color: #e74c3c; font-weight: bold; }
        .activity-list { list-style-type: disc; padding-left: 20px; }
        .quality-metrics { background-color: #e8f5e8; padding: 15px; border-radius: 5px; }
    </style>
</head>
<body>
    <div class="header">
        <h1 class="title">AI原生PBL课程设计</h1>
        <h2 class="subtitle">{{ course.title }}</h2>
    </div>

    <div class="section">
        <h3 class="section-title">课程基本信息</h3>
        <table class="info-table">
            <tr><th>课程名称</th><td>{{ course.title }}</td></tr>
            <tr><th>课程描述</th><td>{{ course.description }}</td></tr>
            <tr><th>教育层级</th><td>{{ course.education_level }}</td></tr>
            <tr><th>适用年级</th><td>{{ course.grade_levels | join(', ') }}年级</td></tr>
            <tr><th>持续时间</th><td>{{ course.duration_weeks }}周 ({{ course.duration_hours }}课时)</td></tr>
            <tr><th>创建时间</th><td>{{ course.created_at[:19] }}</td></tr>
        </table>
    </div>

    <div class="section">
        <h3 class="section-title">学习目标</h3>
        <ol class="objective-list">
            {% for objective in course.learning_objectives %}
            <li>{{ objective }}</li>
            {% endfor %}
        </ol>
    </div>

    <div class="section">
        <h3 class="section-title">项目驱动问题</h3>
        <p><strong>{{ course.driving_question }}</strong></p>
    </div>

    <div class="section">
        <h3 class="section-title">最终产品</h3>
        <ul>
            {% for product in course.final_products %}
            <li>{{ product }}</li>
            {% endfor %}
        </ul>
    </div>

    <div class="section">
        <h3 class="section-title">课程实施阶段</h3>
        {% for phase in course.phases %}
        <div class="phase">
            <div class="phase-title">{{ phase.name }} ({{ phase.duration }})</div>
            <ul class="activity-list">
                {% for activity in phase.activities %}
                <li>{{ activity }}</li>
                {% endfor %}
            </ul>
            <p><strong>推荐AI工具:</strong> {{ phase.ai_tools | join(', ') }}</p>
        </div>
        {% endfor %}
    </div>

    {% if include_assessments and course.assessments %}
    <div class="section">
        <h3 class="section-title">评估体系</h3>
        {% for assessment in course.assessments %}
        <div class="phase">
            <div class="phase-title">{{ assessment.name }} (权重: {{ (assessment.weight * 100) | int }}%)</div>
            <p><strong>类型:</strong> {{ assessment.type }}</p>
            <ul>
                {% for method in assessment.methods %}
                <li>{{ method }}</li>
                {% endfor %}
            </ul>
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if include_resources and course.resources %}
    <div class="section">
        <h3 class="section-title">教学资源</h3>
        {% for resource in course.resources %}
        <div class="phase">
            <div class="phase-title">{{ resource.title }}</div>
            <p><strong>类型:</strong> {{ resource.type }}</p>
            <p>{{ resource.description }}</p>
        </div>
        {% endfor %}
    </div>
    {% endif %}

    {% if course.quality_metrics %}
    <div class="section">
        <h3 class="section-title">课程质量指标</h3>
        <div class="quality-metrics">
            <p><strong>AI能力覆盖度:</strong> {{ (course.quality_metrics.ai_competency_coverage * 100) | int }}%</p>
            <p><strong>PBL方法论完整性:</strong> {{ (course.quality_metrics.pbl_methodology_score * 100) | int }}%</p>
            <p><strong>内容丰富度:</strong> {{ (course.quality_metrics.content_richness * 100) | int }}%</p>
            <p><strong>评估真实性:</strong> {{ (course.quality_metrics.assessment_authenticity * 100) | int }}%</p>
            <p><strong>资源完整性:</strong> {{ (course.quality_metrics.resource_completeness * 100) | int }}%</p>
        </div>
    </div>
    {% endif %}

    <div class="section">
        <h3 class="section-title">设计团队</h3>
        <p>本课程由以下AI专业智能体协作设计完成:</p>
        <ul>
            <li><strong>教育理论专家</strong> - AI时代教育理论和PBL方法论</li>
            <li><strong>课程架构师</strong> - 面向AI时代能力的课程结构设计</li>
            <li><strong>内容设计师</strong> - AI时代场景化学习内容创作</li>
            <li><strong>评估专家</strong> - AI时代核心能力评价体系设计</li>
            <li><strong>素材创作者</strong> - AI时代数字化资源生成</li>
        </ul>
    </div>

    <footer style="margin-top: 50px; text-align: center; color: #7f8c8d; font-size: 12px;">
        <p>Generated by AI-Native PBL Course Design System | {{ datetime.now().strftime('%Y-%m-%d %H:%M:%S') }}</p>
    </footer>
</body>
</html>
        """

        template = Template(html_template)
        html_content = template.render(
            course=course_data,
            include_assessments=include_assessments,
            include_resources=include_resources,
            datetime=datetime
        )

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        return file_path

    async def export_to_json(self, course_data: Dict[str, Any], file_path: Path,
                           include_resources: bool, include_assessments: bool) -> Path:
        """导出为JSON格式"""

        # 创建导出数据的副本
        export_data = course_data.copy()

        # 根据选项过滤内容
        if not include_resources:
            export_data.pop('resources', None)
            export_data.pop('technology_requirements', None)
            export_data.pop('teacher_preparation', None)

        if not include_assessments:
            export_data.pop('assessments', None)

        # 添加导出元数据
        export_data['export_metadata'] = {
            'exported_at': datetime.now().isoformat(),
            'export_format': 'json',
            'includes_resources': include_resources,
            'includes_assessments': include_assessments,
            'generator': 'AI-Native PBL Course Design System'
        }

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

        return file_path

    def get_file_size(self, file_path: Path) -> str:
        """获取文件大小"""
        try:
            size_bytes = file_path.stat().st_size
            if size_bytes < 1024:
                return f"{size_bytes}B"
            elif size_bytes < 1024 * 1024:
                return f"{size_bytes / 1024:.1f}KB"
            else:
                return f"{size_bytes / (1024 * 1024):.1f}MB"
        except:
            return "Unknown"

    def list_exports(self, format_type: Optional[str] = None) -> Dict[str, Any]:
        """列出已导出的文件"""

        exports = []

        if format_type:
            formats = [format_type]
        else:
            formats = ["pdf", "docx", "html", "json"]

        for fmt in formats:
            format_dir = self.export_dir / fmt
            if format_dir.exists():
                for file_path in format_dir.glob(f"*.{fmt}"):
                    exports.append({
                        "format": fmt,
                        "filename": file_path.name,
                        "file_path": str(file_path),
                        "file_size": self.get_file_size(file_path),
                        "created_at": datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
                    })

        return {
            "exports": sorted(exports, key=lambda x: x["created_at"], reverse=True),
            "total_count": len(exports)
        }


# 全局导出服务实例
export_service = CourseExportService()