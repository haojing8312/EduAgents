"""
增强课程导出服务
专门处理增强课程数据格式的文档导出
"""

import os
import json
import uuid
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.platypus.flowables import PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from jinja2 import Template


class EnhancedCourseExportService:
    """增强课程导出服务类"""

    def __init__(self):
        self.export_dir = Path("/home/easegen/EduAgents/backend/exports")
        self.ensure_export_dirs()

    def ensure_export_dirs(self):
        """确保导出目录存在"""
        for format_type in ["pdf", "docx", "html", "json"]:
            (self.export_dir / format_type).mkdir(parents=True, exist_ok=True)

    async def export_enhanced_course(self, course_data: Dict[str, Any], export_format: str,
                                   include_resources: bool = True,
                                   include_assessments: bool = True) -> Dict[str, Any]:
        """
        导出增强课程为指定格式

        Args:
            course_data: 增强课程数据
            export_format: 导出格式 (pdf, docx, html, json)
            include_resources: 是否包含资源
            include_assessments: 是否包含评估
        """

        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        course_id = course_data.get("course_id", str(uuid.uuid4()))
        filename_base = f"enhanced_course_{course_id}_{timestamp}"

        try:
            if export_format == "html":
                return await self._export_html(course_data, filename_base, include_resources, include_assessments)
            elif export_format == "pdf":
                return await self._export_pdf(course_data, filename_base, include_resources, include_assessments)
            elif export_format == "docx":
                return await self._export_docx(course_data, filename_base, include_resources, include_assessments)
            elif export_format == "json":
                return await self._export_json(course_data, filename_base, include_resources, include_assessments)
            else:
                raise ValueError(f"不支持的导出格式: {export_format}")

        except Exception as e:
            raise Exception(f"导出失败: {str(e)}")

    async def _export_html(self, course_data: Dict[str, Any], filename_base: str,
                         include_resources: bool, include_assessments: bool) -> Dict[str, Any]:
        """导出HTML格式"""

        html_template = """
<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{{ course_title }} - AI原生PBL课程设计</title>
    <style>
        body { font-family: "微软雅黑", Arial, sans-serif; margin: 40px; line-height: 1.8; background-color: #fafafa; }
        .container { max-width: 1200px; margin: 0 auto; background: white; padding: 40px; border-radius: 10px; box-shadow: 0 2px 20px rgba(0,0,0,0.1); }
        .header { text-align: center; margin-bottom: 40px; border-bottom: 3px solid #3498db; padding-bottom: 20px; }
        .title { color: #2c3e50; font-size: 32px; margin-bottom: 10px; font-weight: bold; }
        .subtitle { color: #34495e; font-size: 22px; margin-bottom: 5px; }
        .theme { color: #e74c3c; font-size: 18px; font-style: italic; }
        .section { margin-bottom: 35px; }
        .section-title { color: #2980b9; font-size: 20px; border-bottom: 2px solid #3498db; padding-bottom: 8px; margin-bottom: 15px; font-weight: bold; }
        .info-table { width: 100%; border-collapse: collapse; margin: 15px 0; }
        .info-table th, .info-table td { border: 1px solid #ddd; padding: 15px; text-align: left; }
        .info-table th { background-color: #f8f9fa; font-weight: bold; color: #2c3e50; }
        .info-table td { background-color: #ffffff; }
        .objective-list { list-style-type: decimal; padding-left: 25px; }
        .objective-list li { margin-bottom: 8px; padding: 5px; background-color: #f8f9fa; border-radius: 5px; }
        .activity-card { margin: 20px 0; padding: 20px; background-color: #f1f2f6; border-radius: 8px; border-left: 5px solid #e74c3c; }
        .activity-title { color: #e74c3c; font-weight: bold; font-size: 18px; margin-bottom: 10px; }
        .activity-content { margin-bottom: 10px; }
        .activity-steps { margin: 10px 0; padding-left: 20px; }
        .activity-steps li { margin-bottom: 5px; }
        .time-slot { display: flex; margin: 10px 0; padding: 15px; background-color: #e8f5e8; border-radius: 5px; align-items: center; }
        .time-info { font-weight: bold; color: #27ae60; min-width: 200px; }
        .time-activity { color: #2c3e50; }
        .ai-tool-card { margin: 15px 0; padding: 20px; background-color: #fff3cd; border-radius: 8px; border: 1px solid #ffeaa7; }
        .ai-tool-name { color: #d63031; font-weight: bold; font-size: 16px; }
        .teacher-prep { background-color: #e3f2fd; padding: 20px; border-radius: 8px; margin: 15px 0; }
        .prep-category { margin-bottom: 15px; }
        .prep-title { font-weight: bold; color: #1976d2; margin-bottom: 8px; }
        .prep-list { list-style-type: disc; padding-left: 20px; }
        .prep-list li { margin-bottom: 5px; }
        .quality-metrics { background-color: #e8f5e8; padding: 20px; border-radius: 8px; text-align: center; }
        .metric { display: inline-block; margin: 10px; padding: 10px 20px; background-color: #27ae60; color: white; border-radius: 5px; font-weight: bold; }
        .footer { margin-top: 50px; text-align: center; color: #7f8c8d; font-size: 12px; border-top: 1px solid #ddd; padding-top: 20px; }
    </style>
</head>
<body>
    <div class="container">
        <div class="header">
            <h1 class="title">AI原生PBL课程设计</h1>
            <h2 class="subtitle">{{ course_title }}</h2>
            <p class="theme">{{ theme_concept }}</p>
        </div>

        <div class="section">
            <h3 class="section-title">课程基本信息</h3>
            <table class="info-table">
                <tr><th>课程名称</th><td>{{ course_title }}</td></tr>
                <tr><th>主题概念</th><td>{{ theme_concept }}</td></tr>
                <tr><th>机构类型</th><td>{{ institution_type }}</td></tr>
                <tr><th>参与人数</th><td>{{ participants }}人</td></tr>
                <tr><th>目标年龄</th><td>{{ age_group }}</td></tr>
                <tr><th>课程时长</th><td>{{ duration }}小时</td></tr>
                <tr><th>创建时间</th><td>{{ created_at }}</td></tr>
            </table>
        </div>

        <div class="section">
            <h3 class="section-title">学习目标</h3>
            <ol class="objective-list">
                {% for objective in learning_objectives %}
                <li>{{ objective }}</li>
                {% endfor %}
            </ol>
        </div>

        <div class="section">
            <h3 class="section-title">项目驱动问题</h3>
            <p style="font-size: 18px; color: #2c3e50; background-color: #f8f9fa; padding: 15px; border-radius: 5px; border-left: 5px solid #3498db;"><strong>{{ driving_question }}</strong></p>
        </div>

        <div class="section">
            <h3 class="section-title">详细时间安排</h3>
            {% for slot in schedule %}
            <div class="time-slot">
                <div class="time-info">{{ slot.start_time }} - {{ slot.end_time }} ({{ slot.duration_minutes }}分钟)</div>
                <div class="time-activity">{{ slot.activity_type }}</div>
            </div>
            {% endfor %}
        </div>

        <div class="section">
            <h3 class="section-title">详细活动设计</h3>
            {% for activity in detailed_activities %}
            <div class="activity-card">
                <div class="activity-title">{{ activity.title }} ({{ activity.duration_minutes }}分钟)</div>
                <div class="activity-content"><strong>目标：</strong>{{ activity.objective }}</div>
                <div class="activity-content"><strong>所需材料：</strong>{{ activity.materials_needed | join(', ') }}</div>
                <div class="activity-content"><strong>AI工具：</strong>{{ activity.ai_tools_used | join(', ') }}</div>
                <div class="activity-content">
                    <strong>详细步骤：</strong>
                    <ol class="activity-steps">
                        {% for step in activity.step_by_step %}
                        <li>{{ step }}</li>
                        {% endfor %}
                    </ol>
                </div>
                {% if activity.teacher_notes %}
                <div class="activity-content">
                    <strong>教师注意事项：</strong>
                    <ul class="activity-steps">
                        {% for note in activity.teacher_notes %}
                        <li>{{ note }}</li>
                        {% endfor %}
                    </ul>
                </div>
                {% endif %}
            </div>
            {% endfor %}
        </div>

        <div class="section">
            <h3 class="section-title">AI工具使用指导</h3>
            {% for tool in ai_tools_guidance %}
            <div class="ai-tool-card">
                <div class="ai-tool-name">{{ tool.tool_name }}</div>
                <p><strong>描述：</strong>{{ tool.description }}</p>
                <p><strong>使用场景：</strong>{{ tool.use_cases | join(', ') }}</p>
                <div>
                    <strong>使用教程：</strong>
                    <ol style="margin-top: 5px;">
                        {% for step in tool.step_by_step_tutorial %}
                        <li>{{ step }}</li>
                        {% endfor %}
                    </ol>
                </div>
            </div>
            {% endfor %}
        </div>

        <div class="section">
            <h3 class="section-title">教师准备指导</h3>
            <div class="teacher-prep">
                <div class="prep-category">
                    <div class="prep-title">技能要求</div>
                    <ul class="prep-list">
                        {% for skill in teacher_preparation.skill_requirements %}
                        <li>{{ skill }}</li>
                        {% endfor %}
                    </ul>
                </div>

                <div class="prep-category">
                    <div class="prep-title">课前准备</div>
                    <ul class="prep-list">
                        {% for prep in teacher_preparation.pre_course_preparation %}
                        <li>{{ prep }}</li>
                        {% endfor %}
                    </ul>
                </div>

                <div class="prep-category">
                    <div class="prep-title">材料清单</div>
                    <ul class="prep-list">
                        {% for material in teacher_preparation.material_checklist %}
                        <li>{{ material }}</li>
                        {% endfor %}
                    </ul>
                </div>

                <div class="prep-category">
                    <div class="prep-title">应急预案</div>
                    <ul class="prep-list">
                        {% for plan in teacher_preparation.backup_plans %}
                        <li>{{ plan }}</li>
                        {% endfor %}
                    </ul>
                </div>
            </div>
        </div>

        {% if include_assessments %}
        <div class="section">
            <h3 class="section-title">评估体系</h3>
            {% for assessment in assessment_methods %}
            <div class="activity-card">
                <div class="activity-title">{{ assessment.type }} (权重: {{ assessment.weight }})</div>
                <div class="activity-content">
                    <strong>评估方法：</strong>
                    <ul class="activity-steps">
                        {% for method in assessment.methods %}
                        <li>{{ method }}</li>
                        {% endfor %}
                    </ul>
                </div>
            </div>
            {% endfor %}
        </div>
        {% endif %}

        {% if include_resources %}
        <div class="section">
            <h3 class="section-title">所需材料和资源</h3>
            <div style="display: flex; gap: 20px;">
                <div style="flex: 1;">
                    <h4 style="color: #2980b9;">必需材料</h4>
                    <ul class="prep-list">
                        {% for material in required_materials %}
                        <li>{{ material }}</li>
                        {% endfor %}
                    </ul>
                </div>
                <div style="flex: 1;">
                    <h4 style="color: #2980b9;">可选材料</h4>
                    <ul class="prep-list">
                        {% for material in optional_materials %}
                        <li>{{ material }}</li>
                        {% endfor %}
                    </ul>
                </div>
            </div>
        </div>
        {% endif %}

        <div class="section">
            <h3 class="section-title">预期成果</h3>
            <ul class="prep-list">
                {% for outcome in expected_outcomes %}
                <li style="font-size: 16px; font-weight: bold; color: #e74c3c;">{{ outcome }}</li>
                {% endfor %}
            </ul>
        </div>

        <div class="footer">
            <p>本课程由AI原生多智能体PBL课程设计系统生成</p>
            <p>生成时间: {{ created_at }}</p>
            <p>系统版本: Enhanced Course Design System v2.0</p>
        </div>
    </div>
</body>
</html>
        """

        template = Template(html_template)

        # 处理数据
        overview = course_data.get("overview", {})
        teacher_preparation = course_data.get("teacher_preparation", {})

        html_content = template.render(
            course_title=course_data.get("title", "未知课程"),
            theme_concept=course_data.get("theme_concept", ""),
            institution_type=overview.get("institution_type", ""),
            participants=overview.get("participants", 0),
            age_group=overview.get("age_group", ""),
            duration=overview.get("duration", 0),
            created_at=datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            learning_objectives=course_data.get("learning_objectives", []),
            driving_question=course_data.get("driving_question", ""),
            schedule=course_data.get("schedule", []),
            detailed_activities=course_data.get("detailed_activities", []),
            ai_tools_guidance=course_data.get("ai_tools_guidance", []),
            teacher_preparation=teacher_preparation,
            assessment_methods=course_data.get("assessment_methods", []),
            required_materials=course_data.get("required_materials", []),
            optional_materials=course_data.get("optional_materials", []),
            expected_outcomes=course_data.get("expected_outcomes", []),
            include_assessments=include_assessments,
            include_resources=include_resources
        )

        filename = f"{filename_base}.html"
        file_path = self.export_dir / "html" / filename

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_content)

        file_size = os.path.getsize(file_path)

        return {
            "file_name": filename,
            "file_path": str(file_path),
            "file_size": f"{file_size / 1024:.1f}KB",
            "format": "html",
            "created_at": datetime.now().isoformat()
        }

    async def _export_json(self, course_data: Dict[str, Any], filename_base: str,
                         include_resources: bool, include_assessments: bool) -> Dict[str, Any]:
        """导出JSON格式"""

        # 添加导出元数据
        export_data = course_data.copy()
        export_data["export_metadata"] = {
            "exported_at": datetime.now().isoformat(),
            "export_format": "json",
            "includes_resources": include_resources,
            "includes_assessments": include_assessments,
            "generator": "Enhanced Course Design System v2.0"
        }

        filename = f"{filename_base}.json"
        file_path = self.export_dir / "json" / filename

        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(export_data, f, ensure_ascii=False, indent=2)

        file_size = os.path.getsize(file_path)

        return {
            "file_name": filename,
            "file_path": str(file_path),
            "file_size": f"{file_size / 1024:.1f}KB",
            "format": "json",
            "created_at": datetime.now().isoformat()
        }

    async def _export_pdf(self, course_data: Dict[str, Any], filename_base: str,
                        include_resources: bool, include_assessments: bool) -> Dict[str, Any]:
        """导出PDF格式"""

        filename = f"{filename_base}.pdf"
        file_path = self.export_dir / "pdf" / filename

        # 注册中文字体
        try:
            font_paths = [
                "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
                "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                "/System/Library/Fonts/PingFang.ttc"
            ]

            font_registered = False
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('Chinese', font_path))
                        font_registered = True
                        break
                    except:
                        continue

            if not font_registered:
                # 使用Unicode CID字体作为fallback
                from reportlab.pdfbase.cidfonts import UnicodeCIDFont
                pdfmetrics.registerFont(UnicodeCIDFont('STSong-Light'))
                font_name = 'STSong-Light'
            else:
                font_name = 'Chinese'
        except:
            font_name = 'Helvetica'

        doc = SimpleDocTemplate(str(file_path), pagesize=A4)
        styles = getSampleStyleSheet()

        # 自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            spaceAfter=30,
            textColor=colors.darkblue,
            fontName=font_name if font_name != 'Helvetica' else 'Helvetica-Bold'
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12,
            textColor=colors.darkred,
            fontName=font_name if font_name != 'Helvetica' else 'Helvetica-Bold'
        )

        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=11,
            spaceAfter=6,
            fontName=font_name
        )

        story = []

        # 标题
        story.append(Paragraph(f"AI原生PBL课程设计", title_style))
        story.append(Paragraph(f"{course_data.get('title', '未知课程')}", heading_style))
        story.append(Spacer(1, 20))

        # 基本信息
        overview = course_data.get("overview", {})
        basic_info = [
            ["课程名称", course_data.get("title", "")],
            ["主题概念", course_data.get("theme_concept", "")],
            ["机构类型", overview.get("institution_type", "")],
            ["参与人数", f"{overview.get('participants', 0)}人"],
            ["课程时长", f"{overview.get('duration', 0)}小时"],
            ["创建时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S")]
        ]

        info_table = Table(basic_info)
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, -1), font_name),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
            ('BACKGROUND', (1, 0), (1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(Paragraph("课程基本信息", heading_style))
        story.append(info_table)
        story.append(Spacer(1, 20))

        # 学习目标
        story.append(Paragraph("学习目标", heading_style))
        for i, objective in enumerate(course_data.get("learning_objectives", []), 1):
            story.append(Paragraph(f"{i}. {objective}", normal_style))
        story.append(Spacer(1, 15))

        # 项目驱动问题
        story.append(Paragraph("项目驱动问题", heading_style))
        story.append(Paragraph(course_data.get("driving_question", ""), normal_style))
        story.append(Spacer(1, 15))

        # 详细活动
        story.append(Paragraph("详细活动设计", heading_style))
        for activity in course_data.get("detailed_activities", []):
            story.append(Paragraph(f"活动：{activity.get('title', '')} ({activity.get('duration_minutes', 0)}分钟)",
                                 ParagraphStyle('ActivityTitle', parent=normal_style, fontSize=12, textColor=colors.darkred)))
            story.append(Paragraph(f"目标：{activity.get('objective', '')}", normal_style))
            story.append(Paragraph(f"AI工具：{', '.join(activity.get('ai_tools_used', []))}", normal_style))
            story.append(Spacer(1, 10))

        # 教师准备
        teacher_prep = course_data.get("teacher_preparation", {})
        if teacher_prep:
            story.append(PageBreak())
            story.append(Paragraph("教师准备指导", heading_style))

            story.append(Paragraph("课前准备：", ParagraphStyle('SubHeading', parent=normal_style, fontSize=12, textColor=colors.darkblue)))
            for prep in teacher_prep.get("pre_course_preparation", []):
                story.append(Paragraph(f"• {prep}", normal_style))
            story.append(Spacer(1, 10))

        # 生成PDF
        doc.build(story)
        file_size = os.path.getsize(file_path)

        return {
            "file_name": filename,
            "file_path": str(file_path),
            "file_size": f"{file_size / 1024:.1f}KB",
            "format": "pdf",
            "created_at": datetime.now().isoformat()
        }

    async def _export_docx(self, course_data: Dict[str, Any], filename_base: str,
                         include_resources: bool, include_assessments: bool) -> Dict[str, Any]:
        """导出DOCX格式"""

        filename = f"{filename_base}.docx"
        file_path = self.export_dir / "docx" / filename

        doc = Document()

        # 标题
        title = doc.add_heading('AI原生PBL课程设计', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_heading(course_data.get('title', '未知课程'), level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本信息
        doc.add_heading('课程基本信息', level=2)
        overview = course_data.get("overview", {})

        info_table = doc.add_table(rows=6, cols=2)
        info_table.style = 'Table Grid'

        info_data = [
            ("课程名称", course_data.get("title", "")),
            ("主题概念", course_data.get("theme_concept", "")),
            ("机构类型", overview.get("institution_type", "")),
            ("参与人数", f"{overview.get('participants', 0)}人"),
            ("课程时长", f"{overview.get('duration', 0)}小时"),
            ("创建时间", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        ]

        for i, (key, value) in enumerate(info_data):
            info_table.cell(i, 0).text = key
            info_table.cell(i, 1).text = str(value)

        # 学习目标
        doc.add_heading('学习目标', level=2)
        for i, objective in enumerate(course_data.get("learning_objectives", []), 1):
            doc.add_paragraph(f"{i}. {objective}", style='List Number')

        # 项目驱动问题
        doc.add_heading('项目驱动问题', level=2)
        doc.add_paragraph(course_data.get("driving_question", ""))

        # 详细活动
        doc.add_heading('详细活动设计', level=2)
        for activity in course_data.get("detailed_activities", []):
            doc.add_heading(f"{activity.get('title', '')} ({activity.get('duration_minutes', 0)}分钟)", level=3)
            doc.add_paragraph(f"目标：{activity.get('objective', '')}")
            doc.add_paragraph(f"AI工具：{', '.join(activity.get('ai_tools_used', []))}")

            if activity.get('step_by_step'):
                doc.add_paragraph("详细步骤：")
                for step in activity.get('step_by_step', []):
                    doc.add_paragraph(step, style='List Bullet')

        # 教师准备
        teacher_prep = course_data.get("teacher_preparation", {})
        if teacher_prep:
            doc.add_page_break()
            doc.add_heading('教师准备指导', level=2)

            doc.add_heading('课前准备', level=3)
            for prep in teacher_prep.get("pre_course_preparation", []):
                doc.add_paragraph(prep, style='List Bullet')

            doc.add_heading('应急预案', level=3)
            for plan in teacher_prep.get("backup_plans", []):
                doc.add_paragraph(plan, style='List Bullet')

        # 预期成果
        doc.add_heading('预期成果', level=2)
        for outcome in course_data.get("expected_outcomes", []):
            doc.add_paragraph(outcome, style='List Bullet')

        doc.save(str(file_path))
        file_size = os.path.getsize(file_path)

        return {
            "file_name": filename,
            "file_path": str(file_path),
            "file_size": f"{file_size / 1024:.1f}KB",
            "format": "docx",
            "created_at": datetime.now().isoformat()
        }


# 全局增强导出服务实例
enhanced_export_service = EnhancedCourseExportService()